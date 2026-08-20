from __future__ import annotations

import asyncio
import logging
import mimetypes
import re
import shutil
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from bson import ObjectId
from fastapi import BackgroundTasks, UploadFile
from langchain_core.documents import Document

from app.core.config import BACKEND_DIR, settings
from app.db.chroma import get_vectorstore
from app.db.mongodb import get_db
from app.services.retriever_service import invalidate_metadata_cache

logger = logging.getLogger(__name__)
UPLOAD_DIR = BACKEND_DIR / "uploaded_documents"
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}
ALLOWED_MIME_PREFIXES = {
    ".pdf": ("application/pdf",),
    ".docx": ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/zip"),
    ".txt": ("text/plain", "application/octet-stream"),
}


class DocumentProcessingError(RuntimeError):
    pass


class DocumentValidationError(ValueError):
    pass


def documents_collection():
    return get_db()["documents"]


def _utcnow() -> datetime:
    return datetime.now(timezone.utc)


def _clean_text(text: str) -> str:
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _extension(filename: str) -> str:
    return Path(filename or "").suffix.lower()


def _validate_file(upload: UploadFile) -> tuple[str, str]:
    if not upload.filename:
        raise DocumentValidationError("Missing filename.")
    ext = _extension(upload.filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise DocumentValidationError("Only PDF, DOCX, and TXT files are supported.")
    content_type = (upload.content_type or mimetypes.guess_type(upload.filename)[0] or "").lower()
    allowed = ALLOWED_MIME_PREFIXES[ext]
    if content_type and not any(content_type.startswith(item) for item in allowed):
        raise DocumentValidationError("Uploaded file type does not match the allowed document formats.")
    return ext, content_type


def _save_upload_sync(upload: UploadFile, destination: Path, max_bytes: int) -> int:
    destination.parent.mkdir(parents=True, exist_ok=True)
    upload.file.seek(0)
    size = 0
    with destination.open("wb") as output:
        while True:
            chunk = upload.file.read(1024 * 1024)
            if not chunk:
                break
            size += len(chunk)
            if size > max_bytes:
                raise DocumentValidationError(f"File exceeds maximum upload size of {settings.max_upload_mb} MB.")
            output.write(chunk)
    return size


async def upload_document(upload: UploadFile, background_tasks: BackgroundTasks) -> dict[str, Any]:
    ext, _content_type = _validate_file(upload)
    document_id = str(uuid.uuid4())
    stored_name = f"{document_id}{ext}"
    path = UPLOAD_DIR / stored_name
    file_size = await asyncio.to_thread(_save_upload_sync, upload, path, settings.max_upload_mb * 1024 * 1024)
    now = _utcnow()
    record = {
        "_id": document_id,
        "filename": Path(upload.filename).name,
        "file_type": ext.lstrip("."),
        "file_size": file_size,
        "status": "pending",
        "uploaded_at": now,
        "processed_at": None,
        "chunk_count": 0,
        "path": str(path),
        "error": None,
    }
    await documents_collection().insert_one(record)
    logger.info("Document uploaded | id=%s filename=%s size=%s", document_id, record["filename"], file_size)
    background_tasks.add_task(process_document, document_id)
    return {"document_id": document_id, "status": "pending", "message": "Document uploaded. Processing started."}


async def list_documents() -> list[dict[str, Any]]:
    cursor = documents_collection().find({}).sort("uploaded_at", -1)
    return [_serialize_document(item) async for item in cursor]


async def get_document(document_id: str) -> dict[str, Any] | None:
    item = await documents_collection().find_one({"_id": document_id})
    return _serialize_document(item) if item else None


async def delete_document(document_id: str) -> bool:
    item = await documents_collection().find_one({"_id": document_id})
    if not item:
        return False
    await asyncio.to_thread(_delete_vectors_sync, document_id)
    await documents_collection().delete_one({"_id": document_id})
    path = Path(item.get("path") or "")
    if path.exists():
        try:
            path.unlink()
        except OSError:
            logger.warning("Could not remove uploaded file %s", path)
    invalidate_metadata_cache()
    logger.info("Document deleted | id=%s", document_id)
    return True


async def reprocess_document(document_id: str, background_tasks: BackgroundTasks) -> bool:
    item = await documents_collection().find_one({"_id": document_id})
    if not item:
        return False
    background_tasks.add_task(process_document, document_id)
    return True


async def process_document(document_id: str) -> None:
    started = _utcnow()
    await documents_collection().update_one({"_id": document_id}, {"$set": {"status": "processing", "error": None}})
    item = await documents_collection().find_one({"_id": document_id})
    if not item:
        return
    try:
        path = Path(item["path"])
        text, page_map = await asyncio.to_thread(_extract_text, path, item["file_type"])
        text = _clean_text(text)
        if not text:
            raise DocumentProcessingError("No extractable text was found in this document.")
        chunks = _chunk_text(text, settings.chunk_size, settings.chunk_overlap)
        await asyncio.to_thread(_delete_vectors_sync, document_id)
        docs = []
        ids = []
        for index, chunk in enumerate(chunks):
            chunk_id = f"{document_id}:{index}"
            ids.append(chunk_id)
            docs.append(Document(page_content=chunk, metadata={
                "document_id": document_id,
                "filename": item["filename"],
                "file_type": item["file_type"],
                "file_size": item["file_size"],
                "chunk_id": index,
                "page_number": _page_for_offset(page_map, index, len(chunks)),
                "upload_date": item["uploaded_at"].isoformat() if hasattr(item["uploaded_at"], "isoformat") else str(item["uploaded_at"]),
                "source": item["filename"],
            }))
        logger.info("Document chunking | id=%s chunks=%s", document_id, len(docs))
        await asyncio.to_thread(get_vectorstore().add_documents, docs, ids=ids)
        invalidate_metadata_cache()
        await documents_collection().update_one({"_id": document_id}, {"$set": {"status": "completed", "processed_at": _utcnow(), "chunk_count": len(docs), "error": None}})
        logger.info("Document processed | id=%s chunks=%s", document_id, len(docs))
    except Exception as exc:
        await documents_collection().update_one({"_id": document_id}, {"$set": {"status": "failed", "processed_at": started, "error": str(exc)[:1000]}})
        logger.exception("Document processing failed | id=%s", document_id)


def _extract_text(path: Path, file_type: str) -> tuple[str, list[tuple[int, int]]]:
    if file_type == "txt":
        return path.read_text(encoding="utf-8", errors="ignore"), []
    if file_type == "docx":
        from docx import Document as DocxDocument
        doc = DocxDocument(str(path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip()), []
    if file_type == "pdf":
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        parts = []
        page_map = []
        offset = 0
        for page_number, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text() or ""
            if page_text.strip():
                parts.append(page_text)
                offset += len(page_text)
                page_map.append((offset, page_number))
        return "\n".join(parts), page_map
    raise DocumentProcessingError("Unsupported document type.")


def _chunk_text(text: str, chunk_size: int, overlap: int) -> list[str]:
    chunk_size = max(200, chunk_size)
    overlap = max(0, min(overlap, chunk_size // 2))
    chunks = []
    start = 0
    while start < len(text):
        end = min(len(text), start + chunk_size)
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start = end - overlap
    return chunks


def _page_for_offset(page_map: list[tuple[int, int]], index: int, total_chunks: int) -> int | None:
    if not page_map:
        return None
    approx_offset = int((index / max(total_chunks, 1)) * page_map[-1][0])
    for end_offset, page_number in page_map:
        if approx_offset <= end_offset:
            return page_number
    return page_map[-1][1]


def _delete_vectors_sync(document_id: str) -> None:
    try:
        get_vectorstore()._collection.delete(where={"document_id": document_id})
    except Exception:
        logger.exception("Vector deletion failed | document_id=%s", document_id)
        raise


def _serialize_document(item: dict[str, Any]) -> dict[str, Any]:
    return {
        "id": str(item.get("_id")),
        "filename": item.get("filename", ""),
        "file_type": item.get("file_type", ""),
        "file_size": int(item.get("file_size") or 0),
        "status": item.get("status", "pending"),
        "uploaded_at": item.get("uploaded_at"),
        "processed_at": item.get("processed_at"),
        "chunk_count": int(item.get("chunk_count") or 0),
        "error": item.get("error"),
    }
